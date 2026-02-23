import io
import os
import re
from dotenv import load_dotenv
load_dotenv()

import pdfplumber
from pypdf import PdfReader
from docx import Document
from flask import Flask, request, jsonify, render_template
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from jd_parser import parse_job_description

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB

# ── Semantic model removed — using TF-IDF only for stability ──────────────
_semantic_model = None

def _load_semantic_model():
    return None   # sentence-transformers disabled due to tokenizers conflict

# ─────────────────────────────────────────────────────────────────────────────
# UNIVERSAL SKILL DICTIONARIES
# ─────────────────────────────────────────────────────────────────────────────

TECH_SKILLS = {
    # Technology
    "python","javascript","typescript","java","csharp","golang","cpp","c",
    "ruby","swift","kotlin","rust","php","scala","react","vue","angular",
    "html","css","nextjs","svelte","node","flask","django","fastapi","spring",
    "aws","azure","gcp","docker","kubernetes","terraform","linux",
    "microservices","serverless","containerization","sql","mysql","postgresql",
    "mongodb","redis","elasticsearch","tensorflow","pytorch","nlp","ai","llm",
    "rest","graphql","grpc","git","ci","cd","jenkins","kafka","distributed",
    "concurrency","agile","scrum","kanban","jira","confluence",
    # Healthcare
    "patient care","clinical assessment","medication administration","wound care",
    "icu","aged care","triage","ahpra","iv therapy","infection control",
    "emergency care","palliative care","mental health","paediatrics",
    "obstetrics","community nursing","clinical documentation","manual handling",
    # Pharmacy
    "dispensing","compounding","pharmacovigilance","drug interactions",
    "schedule 8","sterile manufacturing","clinical trials","regulatory affairs",
    "quality assurance","gmp","tga","medication management",
    # Medicine
    "diagnosis","clinical reasoning","surgery","general practice",
    "medical imaging","pathology","pharmacology","prescribing",
    # Business / Finance
    "financial analysis","financial modelling","excel","power bi","tableau",
    "stakeholder management","project management","crm","salesforce","marketing",
    "kpi","budgeting","forecasting","accounting","auditing","risk management",
    "business development","supply chain","operations","xero","myob",
    # Law
    "legal research","contract drafting","litigation","compliance",
    "corporate law","conveyancing","family law","criminal law","mediation",
    # Data
    "r","pandas","numpy","jupyter","statistics","data visualisation",
    "machine learning","deep learning","etl","spark","data engineering",
    # Engineering
    "autocad","solidworks","matlab","circuit design","mechanical design","cad","revit",
    # HR
    "talent acquisition","recruitment","onboarding","performance management",
    "hris","payroll","employee relations","workforce planning",
    # Marketing
    "seo","sem","google analytics","figma","photoshop","content marketing",
    "social media","email marketing","brand management","wordpress","hubspot",
    # Hospitality
    "customer service","point of sale","inventory management","food safety",
    "barista","food handling","cash handling",
}

SOFT_SKILLS = {
    "communication","teamwork","collaboration","mentorship","leadership",
    "creativity","curiosity","problemsolving","adaptable","analytical",
    "time management","critical thinking","empathy","attention to detail",
    "organisation","initiative","resilience","multitasking","interpersonal",
    "selfmotivated","work ethic","integrity","accountability",
    "conflict resolution","negotiation","presentation","decision making",
}

EDU_KEYWORDS = {
    "bachelor","master","phd","degree","diploma","certificate","honours",
    "graduate","postgraduate","undergraduate","university","tafe","college",
    "engineering","science","nursing","medicine","law","education","business",
    "computerscience","softwareengineering","informationtechnology",
    "pharmacy","psychology","accounting","finance","marketing","technology",
}

ROLE_KEYWORDS = {
    "graduate","junior","senior","lead","principal","intern","manager",
    "director","coordinator","specialist","analyst","consultant","engineer",
    "developer","nurse","pharmacist","doctor","physician","associate",
    "officer","administrator","advisor","executive","assistant","trainee",
    "apprentice","supervisor","cloud","fullstack","backend","frontend",
    "mobile","embedded","devops","software","student",
}

# Adjacency — partial credit for related skills
SKILL_ADJACENCY = {
    # Tech
    "containerization": {"docker","kubernetes","aws","gcp","azure","terraform","linux"},
    "distributed":      {"aws","gcp","azure","microservices","kafka","node","rest","grpc"},
    "microservices":    {"docker","kubernetes","aws","rest","node","flask","fastapi","grpc"},
    "kubernetes":       {"docker","aws","gcp","azure","containerization"},
    "mlops":            {"python","docker","kubernetes","tensorflow","pytorch"},
    # Healthcare
    "clinical assessment": {"patient care","triage","diagnosis","nursing"},
    "medication administration": {"dispensing","pharmacology","patient care","nursing"},
    # Business
    "financial modelling": {"excel","accounting","financial analysis","forecasting"},
    "stakeholder management": {"project management","communication","leadership"},
}

SYNONYMS = {
    "problem solving":"problemsolving","self-motivated":"selfmotivated",
    "team player":"teamwork","computer science":"computerscience",
    "software engineering":"softwareengineering",
    "information technology":"informationtechnology",
    "c++":"cpp","c#":"csharp","node.js":"node","vue.js":"vue",
    "react.js":"react","ci/cd":"ci","k8s":"kubernetes",
    "rest api":"rest","rest apis":"rest","agile/scrum":"agile",
    "ms excel":"excel","microsoft excel":"excel","version control":"git",
    "node js":"node","javascript":"javascript","scrum":"agile",
    "kanban":"agile","agile methodology":"agile",
}

# Soft skill proxy words — detect soft skills from resume narrative
SOFT_PROXIES = {
    "communication":    {"communicated","communication","customer","service","support","resolved","presented","wrote"},
    "teamwork":         {"collaborated","collaboration","teamwork","team","agile","sprint","cross-functional"},
    "leadership":       {"led","lead","president","coordinator","managed","supervised","mentored","directed"},
    "creativity":       {"designed","built","developed","created","innovative","crafted","architected"},
    "curiosity":        {"learning","explored","researched","projects","personal","self-taught","studied"},
    "mentorship":       {"mentored","taught","trained","supported","guided","coached"},
    "problemsolving":   {"resolved","fixed","debugged","optimised","improved","solutions","troubleshot"},
    "adaptable":        {"adapted","versatile","multiple","cross","flexible","various"},
    "analytical":       {"analysed","analysis","data","evaluated","measured","assessed"},
    "selfmotivated":    {"self","independent","proactive","initiative","driven","motivated"},
    "attention to detail": {"detail","accurate","precision","thorough","careful","quality"},
    "time management":  {"deadline","priorities","schedule","efficient","organised","timely"},
}

# Score weights
W_TECH  = 0.50
W_CTX   = 0.20
W_SOFT  = 0.15
W_ROLE  = 0.10
W_EDU   = 0.05


# ─────────────────────────────────────────────────────────────────────────────
# UTILITIES
# ─────────────────────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    text = text.lower()
    for src, tgt in sorted(SYNONYMS.items(), key=lambda x: -len(x[0])):
        text = text.replace(src, tgt)
    text = re.sub(r'ci\s*/\s*cd', 'ci', text)
    text = re.sub(r'[^a-zA-Z0-9\s]', ' ', text)
    return re.sub(r'\s+', ' ', text).strip()


def tokenize(text: str) -> set:
    """Single + bigram + trigram tokens for multi-word skill matching."""
    tokens = set()
    words = text.split()
    for w in words:
        if len(w) > 1:
            tokens.add(w)
    for i in range(len(words) - 1):
        tokens.add(f"{words[i]} {words[i+1]}")
    for i in range(len(words) - 2):
        tokens.add(f"{words[i]} {words[i+1]} {words[i+2]}")
    return tokens


def extract_experience(text: str) -> int:
    """
    Only extract years of experience when it appears near requirement keywords.
    Avoids pulling company history (e.g. "30+ years of experience") as a job requirement.
    """
    lines = text.lower().split("\n")
    REQUIREMENT_SIGNALS = {
        "require", "minimum", "at least", "must have", "you have",
        "experience required", "years experience", "years of experience",
        "proven experience", "demonstrable", "ideally", "preferred",
        "looking for", "we need", "seeking", "qualifications",
    }
    results = []
    for line in lines:
        # Skip lines that are clearly about the company, not the candidate
        if any(skip in line for skip in [
            "founded", "over 30 years", "30+ years", "25 years", "20 years",
            "company", "our history", "we have been", "in business", "centres",
            "clients", "careers launched", "global", "ftse", "listed",
        ]):
            continue
        # Only extract if line contains a requirement signal
        has_signal = any(signal in line for signal in REQUIREMENT_SIGNALS)
        matches = re.findall(r"(\d+)\s*\+?\s*years?", line)
        if matches and has_signal:
            results.extend(int(m) for m in matches)

    return max(results, default=0)


def tfidf_score(text_a: str, text_b: str) -> float:
    try:
        vecs = TfidfVectorizer(min_df=1, stop_words='english').fit_transform([text_a, text_b])
        return float(cosine_similarity(vecs[0], vecs[1])[0][0])
    except Exception:
        return 0.0


def semantic_score(text_a: str, text_b: str) -> float:
    model = _load_semantic_model()
    if model is None:
        return 0.0
    try:
        from sentence_transformers import util
        embs  = model.encode([text_a, text_b], convert_to_tensor=True)
        return float(max(0.0, min(1.0, util.cos_sim(embs[0], embs[1]))))
    except Exception:
        return 0.0


def context_score(text_a: str, text_b: str) -> float:
    tfidf = tfidf_score(text_a, text_b)
    sem   = semantic_score(text_a, text_b)
    return (0.40 * tfidf + 0.60 * sem) if _load_semantic_model() else tfidf


def smart_skill_overlap(resume_tokens, job_tokens, skill_set):
    required = job_tokens & skill_set
    if not required:
        return 1.0, set(), set(), set()
    exact = set()
    adjacent = set()
    missing = set()
    for skill in required:
        if skill in resume_tokens:
            exact.add(skill)
        elif skill in SKILL_ADJACENCY and (SKILL_ADJACENCY[skill] & resume_tokens):
            adjacent.add(skill)
        else:
            missing.add(skill)
    score = (len(exact) * 1.0 + len(adjacent) * 0.5) / len(required)
    return score, exact, adjacent, missing


def soft_overlap(resume_tokens, job_tokens):
    required = job_tokens & SOFT_SKILLS
    if not required:
        return 1.0, set(), set()
    matched = set()
    missing = set()
    for skill in required:
        proxies = SOFT_PROXIES.get(skill, {skill})
        if skill in resume_tokens or (proxies & resume_tokens):
            matched.add(skill)
        else:
            missing.add(skill)
    return len(matched) / len(required), matched, missing


# ─────────────────────────────────────────────────────────────────────────────
# FILE EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

def extract_pdf(file_bytes: bytes) -> str:
    text = ""
    # Try pdfplumber first (better layout)
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        if text.strip():
            return text.strip()
    except Exception:
        pass
    # Fallback to pypdf
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    except Exception as e:
        raise ValueError(f"Could not read PDF: {e}")
    return text.strip()


def extract_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception as e:
        raise ValueError(f"Could not read DOCX: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────

def analyze_resume(resume: str, job_desc: str):
    resume_clean  = clean_text(resume)
    job_clean     = clean_text(job_desc)
    resume_tokens = tokenize(resume_clean)
    job_tokens    = tokenize(job_clean)

    # 1. Tech / hard skills — exact + adjacency
    tech, exact_tech, adj_tech, miss_tech = smart_skill_overlap(resume_tokens, job_tokens, TECH_SKILLS)

    # 2. Context — TF-IDF + semantic
    ctx = context_score(resume_clean, job_clean)

    # 3. Soft skills — proxy-aware
    soft, matched_soft, miss_soft = soft_overlap(resume_tokens, job_tokens)

    # 4. Role fit — check both cleaned and raw resume
    role_req = job_tokens & ROLE_KEYWORDS
    if not role_req:
        role = 1.0
    else:
        role_matched = resume_tokens & role_req
        for rk in role_req:
            if rk in resume.lower():
                role_matched.add(rk)
        role = min(1.0, len(role_matched) / len(role_req))

    # 5. Education
    edu_req = job_tokens & EDU_KEYWORDS
    if not edu_req:
        edu = 1.0
    else:
        edu_matched = resume_tokens & edu_req
        for ek in edu_req:
            if ek in resume.lower():
                edu_matched.add(ek)
        edu = min(1.0, len(edu_matched) / len(edu_req))

    # 6. Experience penalty
    req_exp  = extract_experience(job_clean)
    user_exp = extract_experience(resume_clean)
    penalty  = 0.0
    if req_exp > 0 and user_exp < req_exp:
        penalty = min(0.20, (req_exp - user_exp) * 0.04)

    # 7. Weighted score
    raw   = W_TECH*tech + W_CTX*ctx + W_SOFT*soft + W_ROLE*role + W_EDU*edu
    score = round(max(0.0, min(100.0, raw * (1 - penalty) * 100)), 1)

    # 8. Label
    if score >= 75:   label = "Strong Match"
    elif score >= 55: label = "Good Match"
    elif score >= 38: label = "Partial Match"
    else:             label = "Weak Match"

    # 9. Suggestions
    truly_missing = sorted(miss_tech)
    msgs = {
        "Strong Match":  "Strong match! Your profile aligns very well with this role.",
        "Good Match":    "Good match — a few additions could make you a top candidate.",
        "Partial Match": "Partial match — you have related experience. Bridge the gaps below to stand out.",
        "Weak Match":    "Significant gaps detected. Focus on the missing skills below.",
    }
    suggestions = [msgs[label]]

    if adj_tech:
        suggestions.append(
            f"Your existing skills partially cover: {', '.join(sorted(adj_tech))}. "
            "Add explicit experience or projects to strengthen these."
        )
    if truly_missing:
        suggestions.append(
            f"No coverage for: {', '.join(truly_missing[:6])}. "
            "Consider courses, certifications, or personal projects."
        )
    if req_exp > 0 and user_exp < req_exp:
        suggestions.append(
            f"Role requires {req_exp}+ years — highlight projects and internships to compensate."
        )
    if miss_soft:
        suggestions.append(
            f"Make these soft skills explicit in your resume: {', '.join(sorted(miss_soft))}."
        )

    breakdown = {
        "tech_skill_score":   round(tech * 100, 1),
        "tfidf_similarity":   round(ctx  * 100, 1),
        "soft_skill_score":   round(soft * 100, 1),
        "role_level_score":   round(role * 100, 1),
        "education_score":    round(edu  * 100, 1),
        "experience_penalty": round(penalty * 100, 1),
    }

    return score, label, truly_missing[:10], suggestions, breakdown


# ─────────────────────────────────────────────────────────────────────────────
# ROUTES
# ─────────────────────────────────────────────────────────────────────────────

@app.route('/')
def home():
    return render_template('index.html')


@app.route('/upload_resume', methods=['POST'])
def upload_resume():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided."}), 400
        file = request.files['file']
        if not file.filename:
            return jsonify({"error": "Empty filename."}), 400
        fname      = file.filename.lower()
        file_bytes = file.read()
        if not file_bytes:
            return jsonify({"error": "File is empty."}), 400
        if fname.endswith('.pdf'):
            text = extract_pdf(file_bytes)
        elif fname.endswith('.docx'):
            text = extract_docx(file_bytes)
        elif fname.endswith('.doc'):
            return jsonify({"error": ".doc not supported — save as .docx and retry."}), 400
        else:
            return jsonify({"error": "Upload a PDF or DOCX file."}), 400
        if not text or len(text.strip()) < 30:
            return jsonify({"error": "Could not extract text — try copy-pasting instead."}), 400
        return jsonify({"text": text.strip(), "char_count": len(text.strip()), "filename": file.filename})
    except ValueError as e:
        return jsonify({"error": str(e)}), 422
    except Exception as e:
        return jsonify({"error": f"Upload failed: {e}"}), 500


@app.route('/parse_jd', methods=['POST'])
def parse_jd():
    try:
        data = request.get_json(silent=True) or {}
        raw  = data.get('raw_jd', '').strip()
        if not raw:
            return jsonify({"error": "raw_jd is required"}), 400
        if len(raw) < 20:
            return jsonify({"error": "Job description too short"}), 400
        return jsonify(parse_job_description(raw))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        data     = request.get_json(silent=True)
        if not data:
            return jsonify({"error": "Invalid JSON"}), 400
        resume   = data.get('resume',   '').strip()
        job_desc = data.get('job_desc', '').strip()
        if not resume or not job_desc:
            return jsonify({"error": "Both resume and job description are required"}), 400
        if len(resume) < 30:
            return jsonify({"error": "Resume too short"}), 400
        score, label, missing, suggestions, breakdown = analyze_resume(resume, job_desc)
        return jsonify({
            "score": score, "match_label": label,
            "missing_keywords": missing, "suggestions": suggestions,
            "breakdown": breakdown,
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ─────────────────────────────────────────────────────────────────────────────
# RUN
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    app.run(debug=True)